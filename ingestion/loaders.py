"""
ingestion/loaders.py
--------------------
Load personal data from multiple sources into a unified format.
Supports: PDF, DOCX, TXT, Markdown, URLs, plain text strings.
"""

import os
import re
import requests
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None


@dataclass
class RawDocument:
    """A raw document loaded from any source, before chunking."""
    content: str
    source: str                        # file path or URL
    source_type: str                   # "pdf", "docx", "txt", "url", "text"
    title: str = ""
    created_at: datetime = field(default_factory=datetime.now)
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        if not self.title:
            self.title = Path(self.source).stem if self.source else "untitled"


# ── PDF ──────────────────────────────────────────────────────────────────────

def load_pdf(path: str) -> RawDocument:
    """Extract text from a PDF file."""
    if PyPDF2 is None:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

    path = str(path)
    text_parts = []

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())

    content = "\n\n".join(text_parts)
    stat = os.stat(path)

    return RawDocument(
        content=content,
        source=path,
        source_type="pdf",
        title=Path(path).stem,
        created_at=datetime.fromtimestamp(stat.st_mtime),
        metadata={"pages": len(reader.pages), "file_size": stat.st_size}
    )


# ── DOCX ─────────────────────────────────────────────────────────────────────

def load_docx(path: str) -> RawDocument:
    """Extract text from a Word document."""
    if DocxDocument is None:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    path = str(path)
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    stat = os.stat(path)

    return RawDocument(
        content=content,
        source=path,
        source_type="docx",
        title=Path(path).stem,
        created_at=datetime.fromtimestamp(stat.st_mtime),
        metadata={"paragraphs": len(paragraphs)}
    )


# ── TXT / MARKDOWN ───────────────────────────────────────────────────────────

def load_text(path: str) -> RawDocument:
    """Load a plain text or markdown file."""
    path = str(path)
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read().strip()

    stat = os.stat(path)
    ext = Path(path).suffix.lower()

    return RawDocument(
        content=content,
        source=path,
        source_type="markdown" if ext in (".md", ".markdown") else "txt",
        title=Path(path).stem,
        created_at=datetime.fromtimestamp(stat.st_mtime),
        metadata={"chars": len(content)}
    )


# ── URL ───────────────────────────────────────────────────────────────────────

def load_url(url: str, timeout: int = 10) -> RawDocument:
    """Scrape readable text from a URL."""
    if BeautifulSoup is None:
        raise ImportError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")

    headers = {"User-Agent": "Mozilla/5.0 (compatible; SecondBrain/1.0)"}
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove noise: scripts, styles, nav, footer
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Try to grab the article/main content first
    main = soup.find("article") or soup.find("main") or soup.find("body")
    text = main.get_text(separator="\n") if main else soup.get_text(separator="\n")

    # Clean up whitespace
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    content = "\n".join(lines)

    # Grab title
    title_tag = soup.find("title")
    title = title_tag.text.strip() if title_tag else url

    return RawDocument(
        content=content,
        source=url,
        source_type="url",
        title=title,
        created_at=datetime.now(),
        metadata={"url": url, "status_code": response.status_code}
    )


# ── RAW TEXT ─────────────────────────────────────────────────────────────────

def load_raw_text(text: str, title: str = "note", source: str = "manual") -> RawDocument:
    """Load a raw text string directly (e.g. typed notes, chat messages)."""
    return RawDocument(
        content=text.strip(),
        source=source,
        source_type="text",
        title=title,
        created_at=datetime.now(),
        metadata={"chars": len(text)}
    )


# ── AUTO LOADER ───────────────────────────────────────────────────────────────

def load_file(path: str) -> RawDocument:
    """Automatically detect file type and load accordingly."""
    path = Path(path)
    ext = path.suffix.lower()

    loaders = {
        ".pdf":      load_pdf,
        ".docx":     load_docx,
        ".doc":      load_docx,
        ".txt":      load_text,
        ".md":       load_text,
        ".markdown": load_text,
        ".rst":      load_text,
    }

    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(loaders.keys())}")

    return loader(str(path))


def load_directory(directory: str, recursive: bool = True) -> list[RawDocument]:
    """Load all supported files from a directory."""
    supported = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown", ".rst"}
    docs = []
    path = Path(directory)

    pattern = "**/*" if recursive else "*"
    for file_path in path.glob(pattern):
        if file_path.suffix.lower() in supported and file_path.is_file():
            try:
                doc = load_file(str(file_path))
                docs.append(doc)
                print(f"  ✓ Loaded: {file_path.name} ({len(doc.content)} chars)")
            except Exception as e:
                print(f"  ✗ Failed: {file_path.name} — {e}")

    return docs
